"""
JobTracker SaaS - Routes Import/Export et Analyse CV
Supporte deux modes:
- Mode Emergent: utilise emergentintegrations (plateforme Emergent)
- Mode Local: utilise les SDKs standards (openai, google-generativeai)
"""

from fastapi import APIRouter, HTTPException, status, Depends, UploadFile, File
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime, timezone
from dotenv import load_dotenv
import os
import json
import csv
import io
import uuid

load_dotenv()

# Try to import emergentintegrations, fallback to standard SDKs
USE_EMERGENT = False
try:
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    USE_EMERGENT = True
except ImportError:
    try:
        from openai import OpenAI
    except ImportError:
        pass

# Import PDF and DOCX extractors
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from utils.auth import get_current_user

router = APIRouter(prefix="/import", tags=["Import"])


def get_db():
    """Dependency injection pour la DB"""
    pass


# Models
class ImportResult(BaseModel):
    success: bool
    imported_count: int
    errors: List[str]
    skipped: int


class CVAnalysisResult(BaseModel):
    score: int  # 0-100
    summary: str
    skills: List[str]
    experience_years: Optional[int]
    strengths: List[str]
    improvements: List[str]
    matching_jobs: List[dict]
    recommendations: str


class ApplicationImport(BaseModel):
    entreprise: str
    poste: str
    type_poste: Optional[str] = "cdi"
    lieu: Optional[str] = None
    moyen: Optional[str] = None
    date_candidature: Optional[str] = None
    lien: Optional[str] = None
    commentaire: Optional[str] = None
    reponse: Optional[str] = "pending"


class DataImportRequest(BaseModel):
    applications: List[dict]


class InterviewImportRequest(BaseModel):
    interviews: List[dict]


# Import interviews from pre-parsed data
@router.post("/interviews/data", response_model=ImportResult)
async def import_interviews_data(
    request: InterviewImportRequest,
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Import interviews from pre-parsed data with duplicate detection"""
    imported = 0
    errors = []
    skipped = 0
    duplicates = 0
    
    # Build a mapping of old candidature_id to new UUID if needed
    # This helps when importing from old system with numeric IDs
    
    for idx, interview_data in enumerate(request.interviews):
        try:
            # Get candidature_id - could be UUID or numeric from old system
            candidature_id = interview_data.get('candidature_id')
            entreprise = interview_data.get('entreprise') or interview_data.get('Entreprise')
            
            # If candidature_id is provided but not a valid UUID, try to find by index/position
            if candidature_id:
                # Check if it's a valid UUID format
                if isinstance(candidature_id, int) or (isinstance(candidature_id, str) and candidature_id.isdigit()):
                    # It's a numeric ID from old system - try to find corresponding application
                    # Get all user applications sorted by date
                    apps = await db.applications.find(
                        {"user_id": current_user["user_id"]},
                        {"_id": 0, "id": 1, "entreprise": 1}
                    ).sort("date_candidature", 1).to_list(1000)
                    
                    numeric_id = int(candidature_id)
                    if 0 < numeric_id <= len(apps):
                        candidature_id = apps[numeric_id - 1]["id"]
                    else:
                        candidature_id = None
            
            # If still no candidature_id, try to find by entreprise
            if not candidature_id and entreprise:
                app = await db.applications.find_one({
                    "user_id": current_user["user_id"],
                    "entreprise": {"$regex": f"^{entreprise}$", "$options": "i"}
                })
                if app:
                    candidature_id = app["id"]
            
            if not candidature_id:
                errors.append(f"Ligne {idx+1}: candidature non trouvée")
                skipped += 1
                continue
            
            # Get date
            date_val = interview_data.get('date_entretien') or interview_data.get('Date Entretien')
            if isinstance(date_val, (int, float)):
                date_str = datetime.fromtimestamp(date_val / 1000, tz=timezone.utc).isoformat()
            elif date_val:
                # Handle format like "2024-11-29 00:00:00.000000"
                date_str = str(date_val).replace(' ', 'T')
                if '.' in date_str and len(date_str.split('.')[-1]) > 6:
                    date_str = date_str[:26]  # Truncate microseconds
            else:
                date_str = datetime.now(timezone.utc).isoformat()
            
            # Check for duplicate (same candidature + similar date)
            existing = await db.interviews.find_one({
                "user_id": current_user["user_id"],
                "candidature_id": candidature_id,
                "date_entretien": {"$regex": f"^{date_str[:10]}"}
            })
            
            if existing:
                duplicates += 1
                continue
            
            # Map statut
            statut_val = interview_data.get('statut') or interview_data.get('Statut') or 'planned'
            statut_str = str(statut_val).lower()
            if 'réalisé' in statut_str or 'realise' in statut_str or 'completed' in statut_str or '✅' in statut_val:
                statut = 'completed'
            elif 'annulé' in statut_str or 'annule' in statut_str or 'cancelled' in statut_str or '❌' in statut_val:
                statut = 'cancelled'
            else:
                statut = 'planned'
            
            interview_doc = {
                "id": str(uuid.uuid4()),
                "user_id": current_user["user_id"],
                "candidature_id": candidature_id,
                "date_entretien": date_str,
                "type_entretien": interview_data.get('type_entretien') or interview_data.get('Type') or 'technical',
                "format_entretien": interview_data.get('format_entretien') or interview_data.get('Format') or 'video',
                "lieu_lien": interview_data.get('lieu_lien') or interview_data.get('lieu_entretien') or interview_data.get('Lieu/Lien') or interview_data.get('Lieu'),
                "interviewer": interview_data.get('interviewer') or interview_data.get('Recruteur'),
                "statut": statut,
                "commentaire": interview_data.get('commentaire') or interview_data.get('Commentaire'),
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
            
            await db.interviews.insert_one(interview_doc)
            imported += 1
            
        except Exception as e:
            errors.append(f"Ligne {idx+1}: {str(e)}")
            skipped += 1
    
    if duplicates > 0:
        errors.insert(0, f"{duplicates} entretien(s) déjà existant(s) ignoré(s)")
    
    return ImportResult(
        success=True,
        imported_count=imported,
        errors=errors[:10],
        skipped=skipped + duplicates
    )


# Import from pre-parsed data (from frontend preview)
@router.post("/data", response_model=ImportResult)
async def import_data(
    request: DataImportRequest,
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Import applications from pre-parsed data with duplicate detection"""
    imported = 0
    errors = []
    skipped = 0
    duplicates = 0
    
    for idx, app_data in enumerate(request.applications):
        try:
            entreprise = app_data.get('entreprise') or app_data.get('Entreprise')
            poste = app_data.get('poste') or app_data.get('Poste')
            
            if not entreprise or not poste:
                errors.append(f"Ligne {idx+1}: entreprise et poste requis")
                skipped += 1
                continue
            
            # Get date
            date_val = app_data.get('date_candidature') or app_data.get('Date (Postulé)')
            if isinstance(date_val, (int, float)):
                date_str = datetime.fromtimestamp(date_val / 1000, tz=timezone.utc).isoformat()
            elif date_val:
                date_str = str(date_val)
            else:
                date_str = datetime.now(timezone.utc).isoformat()
            
            # Check for duplicate (same company + position for same user)
            existing = await db.applications.find_one({
                "user_id": current_user["user_id"],
                "entreprise": {"$regex": f"^{entreprise}$", "$options": "i"},
                "poste": {"$regex": f"^{poste}$", "$options": "i"}
            })
            
            if existing:
                duplicates += 1
                continue
            
            app_doc = {
                "id": str(uuid.uuid4()),
                "user_id": current_user["user_id"],
                "entreprise": entreprise,
                "poste": poste,
                "type_poste": app_data.get('type_poste') or app_data.get('Type') or 'cdi',
                "lieu": app_data.get('lieu') or app_data.get('Lieu'),
                "moyen": app_data.get('moyen') or app_data.get('Moyen'),
                "date_candidature": date_str,
                "lien": app_data.get('lien') or app_data.get('Lien'),
                "commentaire": app_data.get('commentaire') or app_data.get('Commentaire'),
                "reponse": app_data.get('reponse', 'pending'),
                "is_favorite": False,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
            
            await db.applications.insert_one(app_doc)
            imported += 1

            # Handle nested interviews if present
            if 'interviews' in app_data and isinstance(app_data['interviews'], list):
                for interview_data in app_data['interviews']:
                    try:
                        # Normalize date
                        int_date_val = interview_data.get('date_entretien')
                        int_date_str = datetime.now(timezone.utc).isoformat()
                        
                        if isinstance(int_date_val, (int, float)):
                            int_date_str = datetime.fromtimestamp(int_date_val / 1000, tz=timezone.utc).isoformat()
                        elif int_date_val:
                            int_date_str = str(int_date_val).replace(' ', 'T')
                        
                        # Normalize type_entretien to enum values
                        raw_type = str(interview_data.get('type_entretien', 'technical')).lower()
                        type_mapping = {
                            'rh': 'rh', 'hr': 'rh', 'ressources humaines': 'rh',
                            'technique': 'technical', 'technical': 'technical', 'tech': 'technical',
                            'manager': 'manager', 'managerial': 'manager',
                            'final': 'final', 'finale': 'final',
                        }
                        normalized_type = type_mapping.get(raw_type, 'other')
                        
                        # Normalize format_entretien to enum values
                        raw_format = str(interview_data.get('format_entretien', 'video')).lower()
                        format_mapping = {
                            'visio': 'video', 'video': 'video', 'visioconference': 'video', 
                            'visioconférence': 'video', 'teams': 'video', 'zoom': 'video', 'meet': 'video',
                            'téléphone': 'phone', 'telephone': 'phone', 'phone': 'phone', 'tel': 'phone',
                            'présentiel': 'in_person', 'presentiel': 'in_person', 'in_person': 'in_person',
                            'sur site': 'in_person', 'on site': 'in_person', 'onsite': 'in_person',
                        }
                        normalized_format = format_mapping.get(raw_format, 'video')
                        
                        # Create interview linked to this application
                        interview_doc = {
                            "id": str(uuid.uuid4()),
                            "user_id": current_user["user_id"],
                            "candidature_id": app_doc["id"],
                            "date_entretien": int_date_str,
                            "type_entretien": normalized_type,
                            "format_entretien": normalized_format,
                            "lieu_lien": interview_data.get('lieu_lien'),
                            "interviewer": interview_data.get('interviewer'),
                            "statut": interview_data.get('statut', 'planned'),
                            "commentaire": interview_data.get('commentaire'),
                            "created_at": datetime.now(timezone.utc).isoformat(),
                            "updated_at": datetime.now(timezone.utc).isoformat()
                        }
                        
                        await db.interviews.insert_one(interview_doc)
                    except Exception as ie:
                        print(f"Erreur import entretien: {ie}")
                        # Don't fail the whole row for a bad interview
            
        except Exception as e:
            errors.append(f"Ligne {idx+1}: {str(e)}")
            skipped += 1
    
    # Add duplicate info to errors if any
    if duplicates > 0:
        errors.insert(0, f"{duplicates} candidature(s) déjà existante(s) ignorée(s)")
    
    return ImportResult(
        success=True,
        imported_count=imported,
        errors=errors[:10],
        skipped=skipped + duplicates
    )


# ============== AI Helper Functions ==============

async def analyze_cv_with_emergent(api_key: str, user_id: str, cv_text: str, apps_context: str) -> dict:
    """Analyze CV using Emergent integrations"""
    system_message = """Tu es un expert en recrutement et analyse de CV. Tu dois analyser le CV fourni et donner une évaluation complète.

Tu dois retourner ta réponse UNIQUEMENT au format JSON valide suivant (sans texte avant ou après):
{
    "score": <nombre entre 0 et 100>,
    "summary": "<résumé du profil en 2-3 phrases>",
    "skills": ["compétence1", "compétence2", ...],
    "experience_years": <nombre d'années d'expérience estimé ou null>,
    "strengths": ["point fort 1", "point fort 2", ...],
    "improvements": ["amélioration 1", "amélioration 2", ...],
    "matching_jobs": [
        {"title": "Titre du poste", "match_score": <0-100>, "reason": "Raison du match"},
        ...
    ],
    "recommendations": "<conseils détaillés pour améliorer le CV et la recherche d'emploi>"
}

Sois précis et constructif. Le score doit refléter la qualité globale du CV."""

    user_prompt = f"""Analyse ce CV en détail:

--- CONTENU DU CV ---
{cv_text[:4000]}

--- CANDIDATURES RÉCENTES DU CANDIDAT ---
{apps_context}

Retourne l'analyse au format JSON demandé."""

    chat = LlmChat(
        api_key=api_key,
        session_id=f"cv-analysis-{user_id}-{uuid.uuid4().hex[:8]}",
        system_message=system_message
    ).with_model("openai", "gpt-4o")
    
    response = await chat.send_message(UserMessage(text=user_prompt))
    return response


async def analyze_cv_with_openai(api_key: str, cv_text: str, apps_context: str) -> str:
    """Analyze CV using standard OpenAI SDK"""
    system_message = """Tu es un expert en recrutement et analyse de CV. Tu dois analyser le CV fourni et donner une évaluation complète.

Tu dois retourner ta réponse UNIQUEMENT au format JSON valide suivant (sans texte avant ou après):
{
    "score": <nombre entre 0 et 100>,
    "summary": "<résumé du profil en 2-3 phrases>",
    "skills": ["compétence1", "compétence2", ...],
    "experience_years": <nombre d'années d'expérience estimé ou null>,
    "strengths": ["point fort 1", "point fort 2", ...],
    "improvements": ["amélioration 1", "amélioration 2", ...],
    "matching_jobs": [
        {"title": "Titre du poste", "match_score": <0-100>, "reason": "Raison du match"},
        ...
    ],
    "recommendations": "<conseils détaillés pour améliorer le CV et la recherche d'emploi>"
}

Sois précis et constructif. Le score doit refléter la qualité globale du CV."""

    user_prompt = f"""Analyse ce CV en détail:

--- CONTENU DU CV ---
{cv_text[:4000]}

--- CANDIDATURES RÉCENTES DU CANDIDAT ---
{apps_context}

Retourne l'analyse au format JSON demandé."""

    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_prompt}
        ]
    )
    return response.choices[0].message.content


# Import JSON
@router.post("/json", response_model=ImportResult)
async def import_json(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Import applications from JSON file"""
    if not file.filename.endswith('.json'):
        raise HTTPException(status_code=400, detail="Le fichier doit être au format JSON")
    
    try:
        content = await file.read()
        data = json.loads(content.decode('utf-8'))
        
        # Handle both array and object with 'applications' key
        if isinstance(data, dict):
            applications = data.get('applications', data.get('candidatures', []))
        else:
            applications = data
        
        if not isinstance(applications, list):
            raise HTTPException(status_code=400, detail="Format JSON invalide")
        
        imported = 0
        errors = []
        skipped = 0
        
        for idx, app_data in enumerate(applications):
            try:
                # Validate required fields
                if not app_data.get('entreprise') or not app_data.get('poste'):
                    errors.append(f"Ligne {idx+1}: entreprise et poste requis")
                    skipped += 1
                    continue
                
                # Create application
                app_doc = {
                    "id": str(uuid.uuid4()),
                    "user_id": current_user["user_id"],
                    "entreprise": app_data['entreprise'],
                    "poste": app_data['poste'],
                    "type_poste": app_data.get('type_poste', 'cdi'),
                    "lieu": app_data.get('lieu'),
                    "moyen": app_data.get('moyen'),
                    "date_candidature": app_data.get('date_candidature', datetime.now(timezone.utc).isoformat()),
                    "lien": app_data.get('lien'),
                    "commentaire": app_data.get('commentaire'),
                    "reponse": app_data.get('reponse', 'pending'),
                    "is_favorite": app_data.get('is_favorite', False),
                    "created_at": datetime.now(timezone.utc).isoformat(),
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }
                
                await db.applications.insert_one(app_doc)
                imported += 1
                
            except Exception as e:
                errors.append(f"Ligne {idx+1}: {str(e)}")
                skipped += 1
        
        return ImportResult(
            success=True,
            imported_count=imported,
            errors=errors[:10],  # Limit errors
            skipped=skipped
        )
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Fichier JSON invalide")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur d'import: {str(e)}")


# Import CSV
@router.post("/csv", response_model=ImportResult)
async def import_csv(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Import applications from CSV file"""
    if not file.filename.endswith('.csv'):
        raise HTTPException(status_code=400, detail="Le fichier doit être au format CSV")
    
    try:
        content = await file.read()
        text = content.decode('utf-8')
        
        # Parse CSV
        reader = csv.DictReader(io.StringIO(text))
        
        imported = 0
        errors = []
        skipped = 0
        
        # Map possible column names
        field_mapping = {
            'entreprise': ['entreprise', 'company', 'société', 'societe'],
            'poste': ['poste', 'position', 'job', 'titre', 'title'],
            'type_poste': ['type_poste', 'type', 'contract', 'contrat'],
            'lieu': ['lieu', 'location', 'ville', 'city'],
            'moyen': ['moyen', 'source', 'method', 'canal'],
            'date_candidature': ['date_candidature', 'date', 'applied_date'],
            'lien': ['lien', 'link', 'url'],
            'commentaire': ['commentaire', 'comment', 'notes'],
            'reponse': ['reponse', 'status', 'statut', 'response']
        }
        
        def get_field(row, field_name):
            for possible_name in field_mapping.get(field_name, [field_name]):
                if possible_name in row and row[possible_name]:
                    return row[possible_name]
            return None
        
        for idx, row in enumerate(reader):
            try:
                entreprise = get_field(row, 'entreprise')
                poste = get_field(row, 'poste')
                
                if not entreprise or not poste:
                    errors.append(f"Ligne {idx+2}: entreprise et poste requis")
                    skipped += 1
                    continue
                
                app_doc = {
                    "id": str(uuid.uuid4()),
                    "user_id": current_user["user_id"],
                    "entreprise": entreprise,
                    "poste": poste,
                    "type_poste": get_field(row, 'type_poste') or 'cdi',
                    "lieu": get_field(row, 'lieu'),
                    "moyen": get_field(row, 'moyen'),
                    "date_candidature": get_field(row, 'date_candidature') or datetime.now(timezone.utc).isoformat(),
                    "lien": get_field(row, 'lien'),
                    "commentaire": get_field(row, 'commentaire'),
                    "reponse": get_field(row, 'reponse') or 'pending',
                    "is_favorite": False,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }
                
                await db.applications.insert_one(app_doc)
                imported += 1
                
            except Exception as e:
                errors.append(f"Ligne {idx+2}: {str(e)}")
                skipped += 1
        
        return ImportResult(
            success=True,
            imported_count=imported,
            errors=errors[:10],
            skipped=skipped
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur d'import: {str(e)}")


# CV Analysis
@router.post("/analyze-cv", response_model=CVAnalysisResult)
async def analyze_cv(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Analyze CV with AI and provide detailed feedback"""
    # Get API key based on mode
    if USE_EMERGENT:
        api_key = os.environ.get("EMERGENT_LLM_KEY")
    else:
        api_key = os.environ.get("OPENAI_API_KEY")
    
    if not api_key:
        raise HTTPException(status_code=500, detail="AI service not configured. Set OPENAI_API_KEY in .env")
    
    # Check file type
    allowed_types = ['.pdf', '.docx', '.doc', '.txt']
    file_ext = '.' + file.filename.split('.')[-1].lower() if '.' in file.filename else ''
    if file_ext not in allowed_types:
        raise HTTPException(status_code=400, detail=f"Format non supporté. Utilisez: {', '.join(allowed_types)}")
    
    try:
        content = await file.read()
        cv_text = ""
        
        # Extract text based on file type
        if file_ext == '.txt':
            cv_text = content.decode('utf-8')
            
        elif file_ext == '.pdf':
            if not PDF_SUPPORT:
                raise HTTPException(status_code=500, detail="Support PDF non disponible. Installez PyPDF2.")
            try:
                pdf_reader = PdfReader(io.BytesIO(content))
                text_parts = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                cv_text = "\n".join(text_parts)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Erreur lecture PDF: {str(e)}")
                
        elif file_ext in ['.docx', '.doc']:
            if not DOCX_SUPPORT:
                raise HTTPException(status_code=500, detail="Support DOCX non disponible. Installez python-docx.")
            try:
                doc = DocxDocument(io.BytesIO(content))
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                cv_text = "\n".join(text_parts)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Erreur lecture DOCX: {str(e)}")
        
        if not cv_text or len(cv_text.strip()) < 50:
            raise HTTPException(status_code=400, detail="Impossible d'extraire le texte du CV. Le fichier est peut-être vide ou protégé.")
        
        # Get user's applications for matching
        applications = await db.applications.find(
            {"user_id": current_user["user_id"]},
            {"_id": 0, "entreprise": 1, "poste": 1, "type_poste": 1}
        ).limit(10).to_list(10)
        
        apps_context = "\n".join([f"- {a['poste']} chez {a['entreprise']}" for a in applications]) if applications else "Aucune candidature"
        
        # Call AI based on mode
        if USE_EMERGENT:
            response = await analyze_cv_with_emergent(api_key, current_user["user_id"], cv_text, apps_context)
        else:
            response = await analyze_cv_with_openai(api_key, cv_text, apps_context)
        
        # Parse JSON response
        try:
            # Clean response (remove markdown code blocks if present)
            clean_response = response.strip()
            if clean_response.startswith('```'):
                clean_response = clean_response.split('```')[1]
                if clean_response.startswith('json'):
                    clean_response = clean_response[4:]
            clean_response = clean_response.strip()
            
            result = json.loads(clean_response)
            
            # Save analysis to DB
            await db.cv_analyses.insert_one({
                "user_id": current_user["user_id"],
                "filename": file.filename,
                "analysis": result,
                "created_at": datetime.now(timezone.utc).isoformat()
            })
            
            return CVAnalysisResult(
                score=result.get('score', 50),
                summary=result.get('summary', ''),
                skills=result.get('skills', []),
                experience_years=result.get('experience_years'),
                strengths=result.get('strengths', []),
                improvements=result.get('improvements', []),
                matching_jobs=result.get('matching_jobs', []),
                recommendations=result.get('recommendations', '')
            )
            
        except json.JSONDecodeError:
            # If JSON parsing fails, return a structured response anyway
            return CVAnalysisResult(
                score=50,
                summary="Analyse du CV effectuée",
                skills=[],
                experience_years=None,
                strengths=["CV uploadé avec succès"],
                improvements=["Veuillez réessayer pour une analyse détaillée"],
                matching_jobs=[],
                recommendations=response[:500] if response else "Analyse non disponible"
            )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur d'analyse: {str(e)}")


# Get CV analysis history
@router.get("/cv-history")
async def get_cv_history(
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Get user's CV analysis history"""
    analyses = await db.cv_analyses.find(
        {"user_id": current_user["user_id"]},
        {"_id": 0}
    ).sort("created_at", -1).limit(10).to_list(10)
    
    return analyses
